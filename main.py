import requests
import docx

IN_FILE = r'./вопросы.docx'
OUT_FILE = r'./вопрос-ответ.docx'

def GetAnswer(message):
    URL = 'https://llm.api.cloud.yandex.net/foundationModels/v1/completion'
    ID_CATALOG = # <идентификатор_каталога>
    API_KEY = # API_KEY

    headers = {
        'Content-Type': 'application/json',
        'Authorization': API_KEY
    }

    prompt = {
        "modelUri": f"gpt://{ID_CATALOG}/yandexgpt/latest",
        "completionOptions": {
            "stream": False,
            "temperature": 0.6,
            "maxTokens": "2000"
        },
        "messages": [
            {
                "role": "system",
                "text": message
            }
        ]
    }

    response = requests.post(URL, headers=headers, json=prompt)

    if response.status_code != 200:
        return ('Ошибка = ' + str(response.status_code))
    else:
        return response.json()['result']['alternatives'][0]['message']['text']

i = 1
document = docx.Document()
doc = docx.Document(IN_FILE)
for docpara in doc.paragraphs:
    if not docpara.text == '':
        document.add_paragraph(str(i)+') ' + docpara.text)
        document.add_paragraph(GetAnswer(docpara.text))
        document.add_paragraph('-'*20)
        print(f'Ответ на вопрос # {i} ----- {docpara.text}')
        i = i + 1
    else:
        break

document.save(OUT_FILE)